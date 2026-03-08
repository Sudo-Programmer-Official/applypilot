from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

FIXTURE_DIR = Path(__file__).resolve().parent / "resumes"


def generate() -> None:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    _generate_backend_pdf(FIXTURE_DIR / "backend_layout_resume.pdf")
    _generate_compact_docx(FIXTURE_DIR / "compact_resume.docx")


def _generate_backend_pdf(path: Path) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    y = 760

    def line(text: str, font: str = "Helvetica", size: int = 11, x: int = 72, center: bool = False) -> None:
        nonlocal y
        pdf.setFont(font, size)
        draw_x = x
        if center:
            draw_x = (letter[0] - stringWidth(text, font, size)) / 2
        pdf.drawString(draw_x, y, text)
        y -= size + 6

    line("ABHISHEK KUMAR JHA", "Helvetica-Bold", 18, center=True)
    line("Distributed Systems Engineer | Cloud Data Infrastructure & Scalable Platforms", "Helvetica-Oblique", 12, center=True)
    line("+1 (361) 442 9376 | abhishek@example.com | linkedin.com/in/abhishek-jha-ai", "Helvetica", 10, center=True)
    y -= 8
    line("Experience", "Helvetica-Bold", 14)
    line("Braintree Health - Software Engineering Intern May 2025 - Aug 2025", "Helvetica-Bold", 11)
    line("- Built distributed backend services using AWS Lambda and RDS across concurrent workers")
    line("- Improved throughput by redesigning SQL indexing and workload sharding")
    line("Accenture | Software Engineer Jul 2019 - Nov 2021", "Helvetica-Bold", 11)
    line("- Delivered enterprise software modules in large-scale distributed systems")
    y -= 4
    line("Education", "Helvetica-Bold", 14)
    line("Texas A&M University Corpus Christi, MS in Computer Science May 2026 (Expected)")
    y -= 4
    line("Technical Skills", "Helvetica-Bold", 14)
    line("Languages: Python, Java, TypeScript")
    line("Cloud: AWS, Docker, Kubernetes, PostgreSQL, CI/CD")
    y -= 4
    line("Selected Projects", "Helvetica-Bold", 14)
    line("ApplyPilot: ATS-focused resume intelligence system")
    pdf.save()


def _generate_compact_docx(path: Path) -> None:
    doc = Document()

    def add_paragraph(text: str, size: int = 11, bold: bool = False) -> None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold

    add_paragraph("ABHISHEK KUMAR JHA", size=16, bold=True)
    add_paragraph("Backend Engineer")
    add_paragraph("abhishek@example.com | github.com/abhi | linkedin.com/in/abhishek-jha-ai")
    add_paragraph("ProfessionalExperience", size=13, bold=True)
    add_paragraph("Software Engineer @ Example Labs Jan 2022 - Present", bold=True)
    add_paragraph("- Built APIs with FastAPI and PostgreSQL")
    add_paragraph("- Improved deployment reliability through CI/CD automation")
    add_paragraph("TechnicalSkills", size=13, bold=True)
    add_paragraph("Languages: Python, JavaScript")
    add_paragraph("Frameworks: FastAPI, Vue.js")
    add_paragraph("Projects", size=13, bold=True)
    add_paragraph("ApplyPilot: Resume parser and optimization engine")
    add_paragraph("Education", size=13, bold=True)
    add_paragraph("SRM Institute of Science & Technology, B.Tech in Software Engineering May 2019")
    doc.save(path)


if __name__ == "__main__":
    generate()
